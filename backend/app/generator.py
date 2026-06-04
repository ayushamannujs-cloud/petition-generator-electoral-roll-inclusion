"""The petition generator — pours an intake into a .docx document.

This is the deep module: a tiny public interface (`build_petition`) hiding the
layered assembly (scaffold → boilerplate → particulars → standing law).
Section order follows scaffold-rules.md §5; typography follows §1-4.
"""
from __future__ import annotations

import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.intake import PetitionInput

# Gender-driven vocabulary (intake-form.md Q3 — drives every pronoun).
_RELATION = {"male": "son of", "female": "daughter of", "other": "child of"}
_SUBJECT = {"male": "he", "female": "she", "other": "they"}
_POSSESSIVE = {"male": "his", "female": "her", "other": "their"}

# Typography (scaffold-rules.md §2-4).
_FONT = "Bookman Old Style"
_BODY_PT = Pt(12)
_BODY_SPACING = 1.5

_ALIGN = {
    "heading": WD_ALIGN_PARAGRAPH.CENTER,   # structural marker
    "body": WD_ALIGN_PARAGRAPH.JUSTIFY,     # argument
    "party": WD_ALIGN_PARAGRAPH.RIGHT,      # party label
    "left": WD_ALIGN_PARAGRAPH.LEFT,        # district line
}


def build_petition(intake: PetitionInput) -> Document:
    doc = docx.Document()
    _apply_page_geometry(doc)
    _district_line(doc, intake)
    _forum(doc, intake)
    _cause_title(doc, intake)
    _respondents(doc, intake)
    _preamble(doc)
    _facts(doc, intake)
    _relatives(doc, intake)
    _documents(doc, intake)
    _grounds(doc, intake)
    _prayer(doc, intake)
    _closing_tag(doc)
    _declaration(doc, intake)
    _verification(doc, intake)
    _enclosures(doc, intake)
    return doc


# --- typography primitives ---------------------------------------------------

def _apply_page_geometry(doc: Document) -> None:
    sec = doc.sections[0]
    # A4 portrait.
    sec.page_width = Inches(8.2677)   # 210 mm
    sec.page_height = Inches(11.6929)  # 297 mm
    # Asymmetric margins: left & top get the extra half-inch (§1).
    sec.left_margin = Inches(1.5)
    sec.top_margin = Inches(1.5)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)


def _p(doc: Document, text: str, role: str = "body", *, bold: bool = False):
    para = doc.add_paragraph()
    para.alignment = _ALIGN[role]
    if role == "body":
        para.paragraph_format.line_spacing = _BODY_SPACING
    run = para.add_run(text)
    run.font.name = _FONT
    run.font.size = _BODY_PT
    run.bold = bold or role == "heading"
    return para


# --- scaffold sections -------------------------------------------------------

def _district_line(doc: Document, intake: PetitionInput) -> None:
    district = (intake.district or "____").upper()
    _p(doc, f"DISTRICT: {district}", "left", bold=True)


def _forum(doc: Document, intake: PetitionInput) -> None:
    _p(doc, "BEFORE THE HON'BLE ELECTORAL REGISTRATION APPELLATE TRIBUNAL", "heading")


def _cause_title(doc: Document, intake: PetitionInput) -> None:
    epic = intake.epic_no or "____"
    _p(
        doc,
        f"IN THE MATTER OF: an appeal against deletion from the electoral roll "
        f"bearing EPIC No. {epic};",
    )
    _p(doc, "AND", "heading")
    guardian = intake.guardian_name or "____"
    relation = _RELATION[intake.gender]
    address = intake.address or "____"
    _p(
        doc,
        f"IN THE MATTER OF: {intake.appellant_name}, {relation} {guardian}, "
        f"residing at {address}",
    )
    _p(doc, "………. Appellant", "party")


def _respondents(doc: Document, intake: PetitionInput) -> None:
    _p(doc, "VERSUS", "heading")
    _p(doc, "1. The Electoral Registration Officer")
    _p(doc, "2. The District Election Officer")
    _p(doc, "3. The Chief Electoral Officer")
    _p(doc, "………. Respondents", "party")


def _preamble(doc: Document) -> None:
    _p(
        doc,
        "Supplementing the online application preferred by the Appellant, the "
        "Appellant most respectfully submits as under:",
    )
    _p(doc, "MOST RESPECTFULLY SHEWETH:", bold=True)


def _facts(doc: Document, intake: PetitionInput) -> None:
    epic = intake.epic_no or "____"
    address = intake.address or "____"
    ac_name = intake.constituency_name or "____"
    ac_no = intake.constituency_no or "____"
    subject = _SUBJECT[intake.gender]
    _p(
        doc,
        f"That the Appellant is a citizen of India, residing at {address}, and "
        f"is duly entitled to be enrolled as an elector under Article 326 of the "
        f"Constitution of India.",
    )
    _p(
        doc,
        f"That the Appellant was enrolled in the electoral roll of the "
        f"{ac_name} Assembly Constituency (AC No. {ac_no}) bearing EPIC No. "
        f"{epic}, and {subject} has at all material times complied with the "
        f"requirements of the Representation of the People Act, 1950.",
    )
    _beats(doc, intake)
    # Triggered law (4b) — passport unlocks the Passports Act inside the facts.
    if intake.has_passport:
        no = intake.passport_no or "____"
        _p(
            doc,
            f"That the Appellant holds a valid Indian Passport bearing No. {no}, "
            f"which under the Passports Act, 1967 is conclusive proof of the "
            f"Appellant's citizenship and date of birth.",
        )


def _beats(doc: Document, intake: PetitionInput) -> None:
    """Narrative arc — conditional beats in strict chronology (no beat, no line)."""
    if intake.on_draft_roll:
        date = intake.draft_roll_date or "____"
        _p(
            doc,
            f"That the Appellant's name duly appeared in the Draft Electoral Roll "
            f"published on {date}, the Appellant having submitted the enumeration "
            f"form to the Booth Level Officer.",
        )
    if intake.notice_received:
        no = intake.notice_no or "____"
        date = intake.notice_date or "____"
        reason = f" on the ground of {intake.notice_reason}" if intake.notice_reason else ""
        _p(
            doc,
            f"That the Appellant thereafter received a discrepancy/hearing notice "
            f"bearing No. {no} dated {date}{reason}.",
        )
    if intake.hearing_attended:
        officer = intake.hearing_officer or "____"
        date = intake.hearing_date or "____"
        _p(
            doc,
            f"That the Appellant duly attended the hearing before {officer} on "
            f"{date} and submitted all documents in support of the Appellant's "
            f"claim, in the legitimate expectation of inclusion.",
        )
    if intake.under_adjudication:
        date = intake.final_roll_date or "28.02.2026"
        _p(
            doc,
            f"That in the Final Electoral Roll published on {date}, the Appellant's "
            f"name was marked 'Under Adjudication'.",
        )
    if intake.deleted:
        date = intake.deletion_date or "____"
        rider = (
            " without any prior notice or hearing whatsoever"
            if not intake.prior_notice_before_deletion
            else ""
        )
        _p(
            doc,
            f"That the Appellant's name was thereafter deleted vide a Supplementary "
            f"/ Deleted-Electors list dated {date}{rider}, which order is impugned "
            f"herein.",
        )
    if intake.appeal_filed:
        no = intake.appeal_no or "____"
        date = intake.appeal_date or "____"
        _p(
            doc,
            f"That the Appellant has preferred the present online appeal bearing "
            f"No. {no} dated {date}, well within the period of limitation.",
        )


def _held_documents(intake: PetitionInput) -> list[tuple[str, str | None]]:
    """The documents the appellant holds, as (label, number) — feeds table + enclosures."""
    docs: list[tuple[str, str | None]] = []
    if intake.has_voter_id:
        docs.append(("Voter ID / EPIC card", intake.epic_no))
    if intake.has_aadhaar:
        docs.append(("Aadhaar", intake.aadhaar_no))
    if intake.has_passport:
        docs.append(("Passport", intake.passport_no))
    if intake.has_pan:
        docs.append(("PAN", None))
    if intake.has_gst:
        docs.append(("GST registration", None))
    docs.extend((other, None) for other in intake.other_documents)
    return docs


def _documents(doc: Document, intake: PetitionInput) -> None:
    """3a + cascade-enclosure: a 'documents relied upon' table."""
    docs = _held_documents(intake)
    if not docs:
        return
    _p(
        doc,
        "That the Appellant relies upon and produces the following documents in "
        "support of the Appellant's identity, citizenship and ordinary residence:",
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, ("#", "Document", "Number")):
        cell.paragraphs[0].add_run(head).bold = True
    for i, (label, number) in enumerate(docs, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = label
        row[2].text = number or "—"


def _relatives(doc: Document, intake: PetitionInput) -> None:
    """Switch: close relatives on the roll add a citizenship-corroboration block."""
    if not intake.relatives_on_roll:
        return
    _p(
        doc,
        "That the following close relatives of the Appellant are currently "
        "enrolled in the electoral roll, which corroborates the Appellant's "
        "citizenship and ordinary residence within the constituency:",
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, ("Name", "Relationship", "EPIC No.")):
        cell.paragraphs[0].add_run(head).bold = True
    for rel in intake.relatives_on_roll:
        row = table.add_row().cells
        row[0].text = rel.name
        row[1].text = rel.relationship
        row[2].text = rel.epic_no or "—"


def _grounds(doc: Document, intake: PetitionInput) -> None:
    # Standing law (authorities-rules.md §4a) — fixed cluster, always present:
    # PNJ -> Art 14 arbitrariness -> non-application of mind -> right to vote.
    hardship_rider = (
        f" The said deletion has caused acute hardship to the Appellant, who is "
        f"{intake.hardship}."
        if intake.hardship
        else ""
    )
    _p(
        doc,
        "It is submitted that the impugned deletion was effected in breach of the "
        "principles of natural justice (audi alteram partem), the Appellant having "
        "been afforded no notice or hearing prior to the deletion of the "
        "Appellant's name." + hardship_rider,
    )
    _p(
        doc,
        "It is further submitted that the impugned order is arbitrary and violative "
        "of Articles 14 and 21 of the Constitution of India, and suffers from "
        "non-application of mind, being a non-speaking order unsupported by reasons.",
    )
    _p(
        doc,
        "It is humbly submitted that, not being otherwise disqualified under the "
        "Representation of the People Act, 1951, the Appellant ought to have been "
        "retained on the electoral roll, and that the impugned deletion is "
        "unjustified and untenable in law.",
    )
    _p(
        doc,
        "It is submitted that the right to vote and to be enrolled under Article "
        "326 of the Constitution of India is a facet of the Appellant's "
        "constitutional entitlement which the impugned deletion has defeated.",
    )
    # Triggered law (4b): mapped in 2002 SIR + holds Aadhaar -> Motab Shaikh / ADR.
    if intake.mapped_2002_sir and intake.has_aadhaar:
        _p(
            doc,
            "It is submitted that the judgments in Motab Shaikh v. ECI (WP Civil "
            "399/2026) and ADR v. ECI (WP Civil 640/2025) are squarely applicable "
            "to the Appellant herein, as the Appellant's name was mapped in the "
            "2002 SIR and the Appellant holds an Aadhaar.",
        )


def _prayer(doc: Document, intake: PetitionInput) -> None:
    _p(
        doc,
        "The Appellant most humbly prays that this Hon'ble Appellate Tribunal "
        "may graciously be pleased to:",
    )
    reliefs: list[str] = []
    if intake.relief_set_aside:
        reliefs.append("Set aside and/or quash the impugned deletion;")
    if intake.relief_restore:
        reliefs.append("Direct the respondents to restore the Appellant's name on the roll;")
    if intake.relief_interim_vote:
        reliefs.append(
            "Pass an interim order permitting the Appellant to vote in the "
            "ensuing election;"
        )
    # Cascade-prayer: compensation is available only where hardship is pleaded.
    if intake.relief_compensation and intake.hardship:
        reliefs.append(
            "Award compensation for the anguish, mental agony and suffering "
            "caused to the Appellant;"
        )
    # boilerplate-rules.md §6: residuary clause is always the final relief.
    reliefs.append(
        "Pass such further or other order(s) as Your Honour may deem fit and "
        "proper in the interest of justice."
    )
    for i, relief in enumerate(reliefs):
        letter = chr(ord("a") + i)
        _p(doc, f"({letter}) {relief}")


def _closing_tag(doc: Document) -> None:
    _p(doc, "This application is made bona fide and for the ends of justice.")


def _declaration(doc: Document, intake: PetitionInput) -> None:
    _p(doc, "DECLARATION", "heading")
    _p(
        doc,
        "The Appellant declares that the contents of the above application are "
        "true and correct to the best of the Appellant's knowledge and belief, "
        "and the Appellant is aware that making a false declaration is punishable "
        "under Section 31 of the Representation of the People Act, 1950.",
    )


def _verification(doc: Document, intake: PetitionInput) -> None:
    _p(doc, "Place: ____\t\tSignature: ____", "left")
    _p(doc, "Date: ____", "left")


def _enclosures(doc: Document, intake: PetitionInput) -> None:
    _p(doc, "ENCLOSURES:", "left", bold=True)
    docs = _held_documents(intake)
    for i, (label, number) in enumerate(docs, start=1):
        suffix = f" bearing No. {number}" if number else ""
        _p(doc, f"{i}. Copy of {label}{suffix}.", "left")
