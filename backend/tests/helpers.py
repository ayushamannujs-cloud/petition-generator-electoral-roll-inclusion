"""Shared test helpers for reading back generated petitions."""
from __future__ import annotations

from docx.document import Document


def all_text(doc: Document) -> str:
    """Flatten every paragraph (body + tables) into one newline-joined string."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def para_texts(doc: Document) -> list[str]:
    """Ordered list of non-empty body paragraph texts."""
    return [p.text for p in doc.paragraphs if p.text.strip()]


def index_of(doc: Document, needle: str) -> int:
    """Index of the first body paragraph containing needle, else -1."""
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i
    return -1
